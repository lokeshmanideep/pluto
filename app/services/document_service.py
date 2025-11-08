from typing import List, Dict, Any,Tuple
import re
import os
import tempfile
from docx import Document
from docxtpl import DocxTemplate

class DocumentProcessingService:
    
    def __init__(self):
        self.full_document_text = ""
        pass

    def convert_custom_placeholders_to_jinja(self, docx_path: str) -> Tuple[str, List[Dict[str, Any]]]:
        temp_dir = tempfile.mkdtemp()
        temp_docx = os.path.join(temp_dir, "converted_template.docx")
        
        try:
            document = Document(docx_path)
            extracted_placeholders = []
            placeholder_counter = 1
            self.full_document_text = "\n".join([para.text for para in document.paragraphs if para.text.strip()]) 
            # Process each paragraph
            for para in document.paragraphs:
                if not para.text.strip():
                    continue
                full_text = para.text
                placeholders_found = self._find_placeholders_in_text(full_text, placeholder_counter)
                
                if placeholders_found['placeholders']:
                    
                    self._replace_placeholders_in_paragraph(para, placeholders_found['replacements'])
                    extracted_placeholders.extend(placeholders_found['placeholders'])
                    placeholder_counter = placeholders_found['next_counter']
            
            unique_placeholders = self._deduplicate_placeholders(extracted_placeholders)
            document.save(temp_docx)
            
            return temp_docx, unique_placeholders
            
        except Exception as e:
            raise Exception(f"Error converting placeholders: {str(e)}")
    
    def _find_placeholders_in_text(self, text: str, start_counter: int) -> Dict[str, Any]:
        
        placeholders = []
        replacements = []
        counter = start_counter
        
        bracket_pattern = r'\[([A-Za-z0-9 \-_]+)\]'
        blank_pattern = r'\[_{3,}\]'
        
        # Find bracket placeholders
        for match in re.finditer(bracket_pattern, text):
            placeholder_text = match.group(1).strip()
            jinja_name = self._create_jinja_name(placeholder_text)
            if re.match(r'^_+$', placeholder_text):
                continue
            placeholders.append({
                'original': match.group(0),
                'jinja_name': jinja_name,
                'type': 'text',
                'context': self._get_context(text, match.start(), match.end()),
                'description': f"Fill in the value for {placeholder_text}"
            })
            
            replacements.append({
                'original': match.group(0),
                'replacement': f"{{{{ {jinja_name} }}}}"
            })
        
        # Find blank placeholders
        for match in re.finditer(blank_pattern, text):
            jinja_name = f"blank_{counter}"
            
            placeholders.append({
                'original': match.group(0),
                'jinja_name': jinja_name,
                'type': 'text',
                'context': self._get_context(text, match.start(), match.end()),
                'description': f"Fill in blank field #{counter}"
            })
            
            replacements.append({
                'original': match.group(0),
                'replacement': f"{{{{ {jinja_name} }}}}"
            })
            counter += 1
        
        return {
            'placeholders': placeholders,
            'replacements': replacements,
            'next_counter': counter
        }
    
    def _create_jinja_name(self, placeholder_text: str) -> str:
        clean_name = placeholder_text.strip()
        clean_name = re.sub(r'[^a-zA-Z0-9_\s-]', '', clean_name)  # Remove special chars
        clean_name = re.sub(r'[\s-]+', '_', clean_name)  # Replace spaces/hyphens with underscores
        clean_name = re.sub(r'_+', '_', clean_name)  # Remove multiple underscores
        clean_name = clean_name.strip('_')  # Remove leading/trailing underscores
        
        return clean_name if clean_name else 'placeholder'
    
    def _get_context(self, text: str, start: int, end: int, context_length: int = 100) -> str:
        context_start = max(0, start - context_length)
        context_end = min(len(text), end + context_length)
        context = text[context_start:context_end]
        if len(context) == len(text):
            context_start = max(0, self.full_document_text.find(text) - context_length)
            context_end = min(len(self.full_document_text), context_start + context_length * 2)
            context = self.full_document_text[context_start:context_end]
        return context
    
    def _replace_placeholders_in_paragraph(self, paragraph, replacements: List[Dict[str, str]]) -> None:
        if not replacements:
            return
        
        # Get current paragraph text
        current_text = paragraph.text
        
        # Apply all replacements
        for replacement in replacements:
            current_text = current_text.replace(
                replacement['original'], 
                replacement['replacement']
            )
        
        # Clear all runs and add the new text
        for run in paragraph.runs:
            run.clear()
        
        # Add the modified text as a single run
        if current_text.strip():
            paragraph.runs[0].text = current_text
    
    def _deduplicate_placeholders(self, placeholders: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        seen = set()
        unique_placeholders = []
        
        for placeholder in placeholders:
            jinja_name = placeholder['jinja_name']
            if jinja_name not in seen:
                seen.add(jinja_name)
                unique_placeholders.append(placeholder)
        
        return unique_placeholders


    def extract_text_from_docx(self, file_path: str) -> str:
        try:
            doc = Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            return '\n'.join(text_content)
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")
    
    async def process_document(self, file_path: str) -> Tuple[str, List[Dict[str, Any]], str]:
        try:
            if file_path.lower().endswith('.docx'):
                text_content = self.extract_text_from_docx(file_path)
                converted_template_path, all_placeholders = self.convert_custom_placeholders_to_jinja(file_path)
                return text_content, all_placeholders, converted_template_path
            else:
                raise Exception("Unsupported file format. Only .docx files are supported.")
            
        except Exception as e:
            raise Exception(f"Error processing document: {str(e)}")
    
    async def generate_completed_document(self, template_path: str, context: Dict[str, str], document_id: int) -> str:
        try:
            doc = DocxTemplate(template_path)
            for key, value in context.items():
                if key.startswith("blank_"):
                    context[key] = f"__{value}__"
            doc.render(context)
            output_dir = "/tmp/completed_documents"
            os.makedirs(output_dir, exist_ok=True)
            output_path = os.path.join(output_dir, f"completed_document_{document_id}.docx")
            doc.save(output_path)
            return output_path
        except Exception as e:
            raise Exception(f"Error generating completed document: {str(e)}")